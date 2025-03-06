import os
import sys
import json
import tkinter as tk
from tkinter import ttk, messagebox
from docx import Document
import comtypes.client

#########################
# Настраиваем имена файлов
#########################
EMPLOYEES_FILE = "employees.json"    # Лежит рядом с .exe
CONFIG_FILE = "config.json"          # Тоже лежит рядом с .exe

#########################
# 1. Определяем пути
#########################
def get_base_dir():
    """
    Возвращает путь к папке, где лежит .exe (при PyInstaller)
    или .py (при обычном запуске).
    """
    if getattr(sys, 'frozen', False):
        # Запущен как .exe
        return os.path.dirname(sys.executable)
    else:
        # Запущен как .py
        return os.path.dirname(os.path.abspath(__file__))

def get_templates_dir():
    """
    Шаблоны (Postanovlenie.docx, Soprovoditelnoe.docx) лежат в папке 'Templates'
    рядом с .exe (или .py).
    """
    return os.path.join(get_base_dir(), "Templates")

#########################
# 2. Загрузка/Сохранение JSON
#########################
def employees_path():
    """Путь к employees.json рядом с .exe (или .py)."""
    return os.path.join(get_base_dir(), EMPLOYEES_FILE)

def config_path():
    """Путь к config.json рядом с .exe (или .py)."""
    return os.path.join(get_base_dir(), CONFIG_FILE)

def load_employees():
    path = employees_path()
    if not os.path.exists(path):
        data = {
            "employees": [],
            "rank_options": [],
            "position_options": []
        }
        with open(path, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=4)
        return data
    else:
        with open(path, "r", encoding="utf-8") as f:
            return json.load(f)

def save_employees(data):
    path = employees_path()
    with open(path, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=4)

def load_config():
    path = config_path()
    if os.path.exists(path):
        with open(path, "r", encoding="utf-8") as f:
            return json.load(f)
    return {}

def save_config(cfg):
    path = config_path()
    with open(path, "w", encoding="utf-8") as f:
        json.dump(cfg, f, ensure_ascii=False, indent=4)

#########################
# 3. Логика заполнения и печати документов
#########################

def fill_placeholders_run_based(doc: Document, placeholders: dict):
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            for key, val in placeholders.items():
                if key in run.text:
                    run.text = run.text.replace(key, val)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        for key, val in placeholders.items():
                            if key in run.text:
                                run.text = run.text.replace(key, val)

def fill_placeholders(template_in, template_out, placeholders):
    doc = Document(template_in)
    fill_placeholders_run_based(doc, placeholders)
    doc.save(template_out)

def export_to_pdf_and_open(docx_path, pdf_path):
    """Открывает docx в Word, сохраняет в PDF, закрывает Word, затем открывает PDF."""
    word = comtypes.client.CreateObject("Word.Application")
    word.Visible = False
    try:
        doc = word.Documents.Open(docx_path)
        # Убираем печать, оставляем только экспорт в PDF
        doc.ExportAsFixedFormat(pdf_path, 17, OpenAfterExport=False)
        doc.Close(False)
    finally:
        word.Quit()

    # Открываем полученный PDF в стандартной программе просмотра
    os.startfile(pdf_path)


def generate_documents(kusp, date_reg, fio, dob, address, selected_emp_key, employee_dict):
    """
    Заполняет 2 шаблона (Постановление, Сопроводительное), печатает и сохраняет в PDF.
    """
    base_dir = get_base_dir()
    templates_dir = get_templates_dir()

    # Папка Temp для заполненных DOCX
    temp_dir = os.path.join(base_dir, "Temp")
    os.makedirs(temp_dir, exist_ok=True)

    # Папка PDF на рабочем столе
    desktop = os.path.join(os.environ["USERPROFILE"], "Desktop")
    pdf_folder = os.path.join(desktop, "PDF 2025")
    os.makedirs(pdf_folder, exist_ok=True)

    # Пути к шаблонам
    postanovlenie_template = os.path.join(templates_dir, "Postanovlenie.docx")
    soprovoditelnoe_template = os.path.join(templates_dir, "Soprovoditelnoe.docx")

    # Достаём данные сотрудника
    if selected_emp_key in employee_dict:
        emp = employee_dict[selected_emp_key]
    else:
        emp = {"Фамилия": "", "Инициалы": "", "Звание": "", "Должность": "", "Телефон": ""}

    # Формируем словарь плейсхолдеров
    placeholders = {
        "{KUSP}": kusp,
        "{Дата}": date_reg,
        "{FIO}": fio,  # ФИО умершего
        "{ДатаРождения}": dob,
        "{Адрес}": address,
        "{Фамилия}": emp["Фамилия"],
        "{Инициалы}": emp["Инициалы"],
        "{Звание}": emp["Звание"],
        "{Должность}": emp["Должность"],
        "{Телефон}": emp["Телефон"],
    }

    # Итоговые docx/pdf
    postanov_docx = f"КУСП-{kusp} от {date_reg} г. (пост).docx"
    sopr_docx     = f"КУСП-{kusp} от {date_reg} г. (сопр).docx"

    postanov_pdf  = f"КУСП-{kusp} от {date_reg} г. (пост).pdf"
    sopr_pdf      = f"КУСП-{kusp} от {date_reg} г. (сопр).pdf"

    postanov_filled = os.path.join(temp_dir, postanov_docx)
    sopr_filled     = os.path.join(temp_dir, sopr_docx)

    postanov_pdf_out = os.path.join(pdf_folder, postanov_pdf)
    sopr_pdf_out     = os.path.join(pdf_folder, sopr_pdf)

    # Заполняем шаблоны
    fill_placeholders(postanovlenie_template, postanov_filled, placeholders)
    fill_placeholders(soprovoditelnoe_template, sopr_filled, placeholders)

    # Печать + PDF
    export_to_pdf_and_open(postanov_filled, postanov_pdf_out)
    export_to_pdf_and_open(sopr_filled,     sopr_pdf_out)


#########################
# 4. Валидация даты
#########################

def validate_date(new_value):
    for ch in new_value:
        if not (ch.isdigit() or ch == '.'):
            return False
    return True

#########################
# 5. Окно редактирования сотрудников
#########################

def edit_employees_window(root, employees_data, update_combobox_callback):
    top = tk.Toplevel(root)
    top.title("Редактирование списка сотрудников")

    frame_table = tk.Frame(top)
    frame_table.pack(side=tk.TOP, fill=tk.BOTH, expand=True)

    columns = ("Фамилия", "Инициалы", "Звание", "Должность", "Телефон")
    tree = ttk.Treeview(frame_table, columns=columns, show="headings")
    for col in columns:
        tree.heading(col, text=col)
        tree.column(col, width=100)

    scrollbar_y = ttk.Scrollbar(frame_table, orient=tk.VERTICAL, command=tree.yview)
    scrollbar_x = ttk.Scrollbar(frame_table, orient=tk.HORIZONTAL, command=tree.xview)
    tree.configure(yscroll=scrollbar_y.set, xscroll=scrollbar_x.set)

    scrollbar_y.pack(side=tk.RIGHT, fill=tk.Y)
    scrollbar_x.pack(side=tk.BOTTOM, fill=tk.X)
    tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

    # Заполняем таблицу
    for emp in employees_data["employees"]:
        tree.insert("", tk.END, values=(
            emp.get("Фамилия", ""),
            emp.get("Инициалы", ""),
            emp.get("Звание", ""),
            emp.get("Должность", ""),
            emp.get("Телефон", "")
        ))

    frame_buttons = tk.Frame(top)
    frame_buttons.pack(side=tk.BOTTOM, fill=tk.X)

    def add_employee():
        add_win = tk.Toplevel(top)
        add_win.title("Добавить сотрудника")

        tk.Label(add_win, text="Фамилия:").grid(row=0, column=0, padx=5, pady=5, sticky="e")
        ent_surname = tk.Entry(add_win, width=20)
        ent_surname.grid(row=0, column=1, padx=5, pady=5)

        tk.Label(add_win, text="Инициалы:").grid(row=1, column=0, padx=5, pady=5, sticky="e")
        ent_inits = tk.Entry(add_win, width=10)
        ent_inits.grid(row=1, column=1, padx=5, pady=5)

        tk.Label(add_win, text="Звание:").grid(row=2, column=0, padx=5, pady=5, sticky="e")
        ent_rank = tk.Entry(add_win, width=20)
        ent_rank.grid(row=2, column=1, padx=5, pady=5)

        tk.Label(add_win, text="Должность:").grid(row=3, column=0, padx=5, pady=5, sticky="e")
        ent_pos = tk.Entry(add_win, width=30)
        ent_pos.grid(row=3, column=1, padx=5, pady=5)

        tk.Label(add_win, text="Телефон:").grid(row=4, column=0, padx=5, pady=5, sticky="e")
        ent_phone = tk.Entry(add_win, width=15)
        ent_phone.grid(row=4, column=1, padx=5, pady=5)

        def on_save():
            surname = ent_surname.get().strip()
            inits = ent_inits.get().strip()
            rank = ent_rank.get().strip()
            pos = ent_pos.get().strip()
            phone = ent_phone.get().strip()

            if not surname:
                messagebox.showerror("Ошибка", "Фамилия не может быть пустой.")
                return

            new_emp = {
                "Фамилия": surname,
                "Инициалы": inits,
                "Звание": rank,
                "Должность": pos,
                "Телефон": phone
            }
            employees_data["employees"].append(new_emp)
            save_employees(employees_data)  # Сохраняем

            tree.insert("", tk.END, values=(surname, inits, rank, pos, phone))
            add_win.destroy()

        tk.Button(add_win, text="Сохранить", command=on_save).grid(row=5, column=0, columnspan=2, pady=10)

    def delete_employee():
        sel = tree.selection()
        if not sel:
            messagebox.showerror("Ошибка", "Сначала выберите сотрудника.")
            return
        answer = messagebox.askyesno("Подтверждение", "Точно удалить выбранного сотрудника?")
        if not answer:
            return
        item = sel[0]
        values = tree.item(item, "values")

        for emp in employees_data["employees"]:
            if (emp.get("Фамилия") == values[0] and
                emp.get("Инициалы") == values[1] and
                emp.get("Звание") == values[2] and
                emp.get("Должность") == values[3] and
                emp.get("Телефон") == values[4]):
                employees_data["employees"].remove(emp)
                break

        save_employees(employees_data)
        tree.delete(item)

    btn_add = tk.Button(frame_buttons, text="Добавить", command=add_employee)
    btn_add.pack(side=tk.LEFT, padx=5, pady=5)

    btn_del = tk.Button(frame_buttons, text="Удалить", command=delete_employee)
    btn_del.pack(side=tk.LEFT, padx=5, pady=5)

    def on_close():
        update_combobox_callback()
        top.destroy()

    top.protocol("WM_DELETE_WINDOW", on_close)


#########################
# 6. Основное окно
#########################

def main():
    employees_data = load_employees()
    config_data = load_config()

    employee_dict = {
        f"{emp['Фамилия']} {emp['Инициалы']}": emp for emp in employees_data["employees"]
    }

    root = tk.Tk()
    root.title("Документы по трупу 15 отдела полиции     v. 1.0")

    vcmd_date = (root.register(validate_date), "%P")

    def update_employee_combobox():
        updated_data = load_employees()
        employee_dict.clear()
        for emp in updated_data["employees"]:
            key = f"{emp['Фамилия']} {emp['Инициалы']}"
            employee_dict[key] = emp
        combo_employee["values"] = list(employee_dict.keys())

    def on_button_click():
        kusp_val = entry_kusp.get().strip()
        date_val = entry_date_reg.get().strip()
        fio_val  = entry_fio.get().strip()
        dob_val  = entry_dob.get().strip()
        addr_val = "Санкт-Петербург, " + entry_address.get().strip()

        try:
            generate_documents(kusp_val, date_val, fio_val, dob_val,
                               addr_val, combo_employee.get(), employee_dict)
            messagebox.showinfo(
                "Успех",
                "Документы сохранены в папке PDF 2025 на рабочем столе!"
            )
        except Exception as e:
            messagebox.showerror("Ошибка", f"Произошла ошибка: {e}")

    def on_edit_button():
        edit_employees_window(root, employees_data, update_employee_combobox)

    def on_close():
        cfg = {"selected_employee": combo_employee.get()}
        save_config(cfg)
        root.destroy()

    # --- GUI ELEMENTS ---
    frm_emp = tk.Frame(root)
    frm_emp.grid(row=0, column=0, columnspan=2, sticky="w", pady=5)

    tk.Label(frm_emp, text="Сотрудник:").pack(side=tk.LEFT, padx=5)
    combo_employee = ttk.Combobox(frm_emp, values=list(employee_dict.keys()), width=30)
    combo_employee.pack(side=tk.LEFT, padx=5)

    btn_edit = tk.Button(frm_emp, text="Изменить", command=on_edit_button)
    btn_edit.pack(side=tk.LEFT, padx=10)

    label_kusp = tk.Label(root, text="КУСП:")
    label_kusp.grid(row=1, column=0, padx=5, pady=5, sticky="w")
    entry_kusp = tk.Entry(root, width=20)
    entry_kusp.grid(row=1, column=1, padx=5, pady=5, sticky="w")

    label_date_reg = tk.Label(root, text="Дата регистрации:")
    label_date_reg.grid(row=2, column=0, padx=5, pady=5, sticky="w")
    entry_date_reg = tk.Entry(root, width=20, validate="key", validatecommand=vcmd_date)
    entry_date_reg.grid(row=2, column=1, padx=5, pady=5, sticky="w")

    label_fio = tk.Label(root, text="ФИО умершего (род. падеж):")
    label_fio.grid(row=3, column=0, padx=5, pady=5, sticky="w")
    entry_fio = tk.Entry(root, width=50)
    entry_fio.grid(row=3, column=1, padx=5, pady=5, sticky="w")

    label_dob = tk.Label(root, text="Дата рождения:")
    label_dob.grid(row=4, column=0, padx=5, pady=5, sticky="w")
    entry_dob = tk.Entry(root, width=20, validate="key", validatecommand=vcmd_date)
    entry_dob.grid(row=4, column=1, padx=5, pady=5, sticky="w")

    label_address = tk.Label(root, text="Адрес (без 'Санкт-Петербург, '):")
    label_address.grid(row=5, column=0, padx=5, pady=5, sticky="w")
    entry_address = tk.Entry(root, width=50)
    entry_address.grid(row=5, column=1, padx=5, pady=5, sticky="w")

    button_generate = tk.Button(root, text="Сохранить в PDF и открыть", command=on_button_click)
    button_generate.grid(row=6, column=0, columnspan=2, padx=5, pady=10, sticky="w")

    # Восстанавливаем выбранного сотрудника
    if "selected_employee" in config_data:
        combo_employee.set(config_data["selected_employee"])

    root.protocol("WM_DELETE_WINDOW", on_close)
    root.mainloop()

if __name__ == "__main__":
    main()
